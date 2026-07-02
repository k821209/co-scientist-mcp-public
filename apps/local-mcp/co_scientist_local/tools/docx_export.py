"""Render a markdown manuscript to a native .docx via python-docx.

Used for non-"paper" document types (reports / other). Pandoc's OOXML output
crashes Hancom Office (한컴오피스) — see the soffice round-trip workaround in
exports.py (dev-todo P0-1). python-docx writes a plain, native package that
Hancom opens cleanly, so reports skip pandoc entirely.

Scope: full CommonMark + GFM tables (markdown-it-py "commonmark" preset with
the table rule enabled). We deliberately avoid the "gfm-like" preset because
it enables linkify, which errors without linkify-it-py installed.

Supported blocks:  headings h1-h6, paragraphs, bullet/ordered lists (nested),
                   GFM tables, blockquotes, fenced/indented code, thematic
                   breaks, images.
Supported inline:  bold, italic, inline code, links (real hyperlinks),
                   images, soft/hard breaks.

Images use the same staged-file convention as the pandoc path: a body embed
of `![alt](figure_N.png)` resolves against `asset_dir` (the export temp dir
where figure blobs were written).
"""
from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.tree import SyntaxTreeNode

# Cap image width to the usable text column of a default Letter page
# (8.5in − 1in margins each side). Smaller images keep their native size.
_MAX_IMAGE_WIDTH = Inches(6.0)
_MONO_FONT = "Consolas"
_LINK_COLOR = RGBColor(0x05, 0x63, 0xC1)

# Report design template. Reports/other docs open in Hancom Office, so we pick a
# CJK-safe default font and style headings/tables for a polished, readable look.
# The East Asian face drives Hangul glyphs; the Latin face drives ASCII.
_BODY_FONT_EASTASIA = "맑은 고딕"          # Malgun Gothic — ships with Korean Windows/Hancom
_BODY_FONT_LATIN = "Times New Roman"        # default export body face (Latin/ASCII)
_BODY_FONT_SIZE = Pt(10.5)
_BODY_LINE_SPACING = 1.15
_HEADING_COLOR = RGBColor(0x1F, 0x38, 0x64)  # deep navy — reads as a section title
# Built-in table style confirmed present in python-docx's default template;
# gives header shading + banded rows out of the box.
_TABLE_STYLE = "Light Grid Accent 1"
# (level, size pt, space-before pt, space-after pt)
_HEADING_SPECS = [(1, 16, 12, 6), (2, 13, 10, 4), (3, 11.5, 8, 3)]


def _set_eastasia_font(style, eastasia: str) -> None:
    """Set a style's East Asian font face via oxml.

    python-docx's `.font.name` only writes the `w:ascii`/`w:hAnsi` attributes;
    Hangul glyphs are selected by `w:eastAsia`, which has no high-level API.
    """
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), eastasia)


def _set_doc_default_font(doc, latin: str, eastasia: str) -> None:
    """Set the document-wide default run font (`w:docDefaults`), so EVERY
    style/paragraph that doesn't override it inherits this face. This is the
    most reliable way to force a base font across a whole docx (incl. one
    pandoc generates from this file as its `--reference-doc`)."""
    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn("w:docDefaults"))
    if docDefaults is None:
        docDefaults = styles_el.makeelement(qn("w:docDefaults"), {})
        styles_el.insert(0, docDefaults)
    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = docDefaults.makeelement(qn("w:rPrDefault"), {})
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        rPr = rPrDefault.makeelement(qn("w:rPr"), {})
        rPrDefault.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), latin)
    rFonts.set(qn("w:eastAsia"), eastasia)


def apply_base_font_to_docx(path) -> None:
    """Swap an existing .docx's base font to the export font (Times New Roman)
    at 1.15 line spacing, WITHOUT touching any other style — so a pandoc-
    generated doc keeps its "Table" style (borders/shading), headings, etc.,
    and only the default typeface changes. Sets the document-wide default font
    (so table cells + body inherit it) plus the Normal style's font + spacing.
    """
    doc = Document(str(path))
    _set_doc_default_font(doc, _BODY_FONT_LATIN, _BODY_FONT_EASTASIA)
    try:
        normal = doc.styles["Normal"]
        normal.font.name = _BODY_FONT_LATIN
        _set_eastasia_font(normal, _BODY_FONT_EASTASIA)
        normal.paragraph_format.line_spacing = _BODY_LINE_SPACING
    except KeyError:
        pass
    # Pandoc's figure-caption styles default to italic; drop it so captions
    # read as regular text (the bold "Figure N." label comes from the alt
    # markdown, and any genuine *italics* keep their direct run formatting).
    for style_name in ("Image Caption", "Caption"):
        try:
            doc.styles[style_name].font.italic = False
        except KeyError:
            pass
    doc.save(str(path))


def _apply_report_template(doc) -> None:
    """Style the document's Normal + Heading styles for report/other exports.

    Sets a CJK-friendly body font (Latin + East Asian faces), comfortable body
    size with relaxed line spacing, and navy headings with breathing room. All
    block renderers inherit from these styles, so this colors the whole doc.
    """
    normal = doc.styles["Normal"]
    normal.font.name = _BODY_FONT_LATIN
    normal.font.size = _BODY_FONT_SIZE
    _set_eastasia_font(normal, _BODY_FONT_EASTASIA)
    pf = normal.paragraph_format
    pf.line_spacing = _BODY_LINE_SPACING
    pf.space_after = Pt(6)

    for level, size, before, after in _HEADING_SPECS:
        style = doc.styles[f"Heading {level}"]
        style.font.name = _BODY_FONT_LATIN
        style.font.size = Pt(size)
        style.font.color.rgb = _HEADING_COLOR
        _set_eastasia_font(style, _BODY_FONT_EASTASIA)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def _md() -> MarkdownIt:
    # "commonmark" + table only — NOT "gfm-like" (that enables linkify, which
    # raises ModuleNotFoundError without linkify-it-py).
    return MarkdownIt("commonmark").enable("table")


def _add_hyperlink(paragraph, url: str, runs_spec: list[tuple[str, bool, bool, bool]]):
    """Append a real Word hyperlink to `paragraph`.

    python-docx has no hyperlink API, so we build the `<w:hyperlink>` element
    by hand: register the URL as an external relationship, then nest styled
    runs inside the hyperlink wrapper. `runs_spec` is a list of
    (text, bold, italic, code) tuples so a link can carry mixed formatting.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    for text, bold, italic, code in runs_spec:
        r = paragraph._p.makeelement(qn("w:r"), {})
        rpr = r.makeelement(qn("w:rPr"), {})
        # Hyperlink-blue + underline so it reads as a link in any viewer.
        color = rpr.makeelement(qn("w:color"), {qn("w:val"): "0563C1"})
        rpr.append(color)
        u = rpr.makeelement(qn("w:u"), {qn("w:val"): "single"})
        rpr.append(u)
        if bold:
            rpr.append(rpr.makeelement(qn("w:b"), {}))
        if italic:
            rpr.append(rpr.makeelement(qn("w:i"), {}))
        if code:
            rfonts = rpr.makeelement(qn("w:rFonts"), {qn("w:ascii"): _MONO_FONT,
                                                      qn("w:hAnsi"): _MONO_FONT})
            rpr.append(rfonts)
        r.append(rpr)
        t = r.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
        t.text = text
        r.append(t)
        hyperlink.append(r)
    paragraph._p.append(hyperlink)


def _styled_run(paragraph, text: str, *, bold: bool, italic: bool, code: bool):
    run = paragraph.add_run(text)
    run.bold = bold or None
    run.italic = italic or None
    if code:
        run.font.name = _MONO_FONT
    return run


def _add_image(paragraph, path: pathlib.Path):
    run = paragraph.add_run()
    shape = run.add_picture(str(path))
    if shape.width > _MAX_IMAGE_WIDTH:
        ratio = _MAX_IMAGE_WIDTH / shape.width
        shape.width = _MAX_IMAGE_WIDTH
        shape.height = int(shape.height * ratio)
    return run


def _fix_drawing_ids(doc) -> None:
    """Give every embedded drawing a unique, non-zero visual id.

    python-docx hardcodes each picture's `pic:cNvPr id="0"`. Word and
    LibreOffice tolerate the zero id, but Hancom Office (한컴오피스) refuses to
    draw a picture whose visual id is 0 — it shows a broken-image box even
    though the package is valid (EXP-2). Renumber both `wp:docPr` and
    `pic:cNvPr` across the document to unique ids starting at 1.
    """
    n = 0
    for el in doc.element.iter():
        tag = el.tag
        if tag.endswith("}docPr") or tag.endswith("}cNvPr"):
            n += 1
            el.set("id", str(n))


def _collect_inline_runs(
    node: SyntaxTreeNode, bold: bool, italic: bool, code: bool
) -> list[tuple[str, bool, bool, bool]]:
    """Flatten an inline subtree into (text, bold, italic, code) run specs.

    Used for link contents, where everything must live inside the single
    `<w:hyperlink>` wrapper rather than as sibling paragraph runs.
    """
    out: list[tuple[str, bool, bool, bool]] = []
    for child in node.children:
        if child.type == "text":
            out.append((child.content, bold, italic, code))
        elif child.type == "code_inline":
            out.append((child.content, bold, italic, True))
        elif child.type == "strong":
            out.extend(_collect_inline_runs(child, True, italic, code))
        elif child.type == "em":
            out.extend(_collect_inline_runs(child, bold, True, code))
        elif child.type in ("softbreak", "hardbreak"):
            out.append((" ", bold, italic, code))
        else:
            out.extend(_collect_inline_runs(child, bold, italic, code))
    return out


def _render_inline(
    paragraph, node: SyntaxTreeNode, asset_dir: pathlib.Path | None,
    *, bold: bool = False, italic: bool = False, code: bool = False,
):
    """Render an inline node's children into runs on `paragraph`."""
    for child in node.children:
        t = child.type
        if t == "text":
            _styled_run(paragraph, child.content, bold=bold, italic=italic, code=code)
        elif t == "code_inline":
            _styled_run(paragraph, child.content, bold=bold, italic=italic, code=True)
        elif t == "strong":
            _render_inline(paragraph, child, asset_dir, bold=True, italic=italic, code=code)
        elif t == "em":
            _render_inline(paragraph, child, asset_dir, bold=bold, italic=True, code=code)
        elif t == "link":
            url = child.attrs.get("href", "")
            runs = _collect_inline_runs(child, bold, italic, code)
            if url and runs:
                _add_hyperlink(paragraph, str(url), runs)
            else:
                _render_inline(paragraph, child, asset_dir, bold=bold, italic=italic, code=code)
        elif t == "image":
            src = str(child.attrs.get("src", ""))
            resolved = _resolve_image(src, asset_dir)
            if resolved is not None:
                _add_image(paragraph, resolved)
            else:
                # Unresolved image → fall back to its alt text as plain run.
                alt = child.content or src
                if alt:
                    _styled_run(paragraph, alt, bold=bold, italic=italic, code=code)
        elif t == "softbreak":
            _styled_run(paragraph, " ", bold=bold, italic=italic, code=code)
        elif t == "hardbreak":
            paragraph.add_run().add_break()


def _resolve_image(src: str, asset_dir: pathlib.Path | None) -> pathlib.Path | None:
    """Resolve an image src to a local file, if it points at a staged asset.

    Only local basenames staged in `asset_dir` are embedded; remote URLs and
    missing files return None so the caller can fall back to alt text.
    """
    if not src or asset_dir is None:
        return None
    if "://" in src:  # remote URL — python-docx can't fetch it
        return None
    candidate = asset_dir / pathlib.Path(src).name
    return candidate if candidate.is_file() else None


def _render_block(doc, node: SyntaxTreeNode, asset_dir: pathlib.Path | None,
                  list_level: int = 0):
    """Render a single block-level node into the document."""
    t = node.type
    if t == "heading":
        level = int(node.tag[1])  # h1..h6 → 1..6
        heading = doc.add_heading(level=min(level, 6))
        inline = node.children[0] if node.children else None
        if inline is not None:
            _render_inline(heading, inline, asset_dir)
    elif t == "paragraph":
        inline = node.children[0] if node.children else None
        # A paragraph that is a single image becomes a centered figure with a
        # caption underneath (the image's alt text). The caption renders its
        # markdown — so a "**Figure N.**" label is bold and the rest regular,
        # while genuine *italics* (e.g. species names) are preserved — instead
        # of force-italicizing the whole line.
        if inline is not None and _is_lone_image(inline):
            img = inline.children[0]
            resolved = _resolve_image(str(img.attrs.get("src", "")), asset_dir)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if resolved is not None:
                _add_image(p, resolved)
                caption = img.content or img.attrs.get("alt") or ""
                if caption:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if img.children:
                        _render_inline(cap, img, asset_dir)
                    else:
                        cap.add_run(caption)
                    for run in cap.runs:
                        run.font.size = Pt(9)
            else:
                run = p.add_run(img.content or str(img.attrs.get("src", "")))
                run.italic = True
            return
        p = doc.add_paragraph()
        if inline is not None:
            _render_inline(p, inline, asset_dir)
    elif t in ("bullet_list", "ordered_list"):
        style = "List Bullet" if t == "bullet_list" else "List Number"
        if list_level > 0:
            style = f"{style} {min(list_level + 1, 3)}"
        for item in node.children:
            if item.type != "list_item":
                continue
            _render_list_item(doc, item, asset_dir, style, list_level)
    elif t == "blockquote":
        for child in node.children:
            if child.type == "paragraph":
                inline = child.children[0] if child.children else None
                p = doc.add_paragraph(style="Quote")
                if inline is not None:
                    _render_inline(p, inline, asset_dir)
            else:
                _render_block(doc, child, asset_dir, list_level)
    elif t in ("code_block", "fence"):
        p = doc.add_paragraph()
        run = p.add_run(node.content.rstrip("\n"))
        run.font.name = _MONO_FONT
        run.font.size = Pt(9)
    elif t == "hr":
        doc.add_paragraph("─" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif t == "table":
        _render_table(doc, node, asset_dir)


def _is_lone_image(inline: SyntaxTreeNode) -> bool:
    kids = [c for c in inline.children if c.type != "softbreak"]
    return len(kids) == 1 and kids[0].type == "image"


def _render_list_item(doc, item: SyntaxTreeNode, asset_dir, style: str, list_level: int):
    for child in item.children:
        if child.type == "paragraph":
            inline = child.children[0] if child.children else None
            p = doc.add_paragraph(style=style)
            if inline is not None:
                _render_inline(p, inline, asset_dir)
        elif child.type in ("bullet_list", "ordered_list"):
            _render_block(doc, child, asset_dir, list_level + 1)
        else:
            _render_block(doc, child, asset_dir, list_level)


def _set_table_look(table, *, header: bool, banded_rows: bool) -> None:
    """Set the table's `<w:tblLook>` so the style's conditional formatting
    (header shading, row banding) actually renders. python-docx exposes no
    high-level API for this, so we write the element attributes directly.
    """
    tbl_pr = table._tbl.tblPr
    look = tbl_pr.find(qn("w:tblLook"))
    if look is None:
        look = tbl_pr.makeelement(qn("w:tblLook"), {})
        tbl_pr.append(look)
    look.set(qn("w:firstRow"), "1" if header else "0")
    look.set(qn("w:lastRow"), "0")
    look.set(qn("w:firstColumn"), "0")
    look.set(qn("w:lastColumn"), "0")
    look.set(qn("w:noHBand"), "0" if banded_rows else "1")
    look.set(qn("w:noVBand"), "1")


def _render_table(doc, node: SyntaxTreeNode, asset_dir):
    rows: list[list[SyntaxTreeNode]] = []
    for section in node.children:  # thead / tbody
        for tr in section.children:
            if tr.type == "tr":
                rows.append([c for c in tr.children if c.type in ("th", "td")])
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = _TABLE_STYLE
    # Turn on the style's conditional formatting: header-row + banded rows,
    # no first/last-column emphasis (so data columns read evenly). Without a
    # `<w:tblLook>` the banding/header shading don't render in some viewers.
    _set_table_look(table, header=True, banded_rows=True)
    for ri, cells in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].text = ""  # clear default empty run
            if ci < len(cells):
                inline = cells[ci].children[0] if cells[ci].children else None
                if inline is not None:
                    _render_inline(cell.paragraphs[0], inline, asset_dir)
            if ri == 0:  # bold header row
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def render_markdown_to_docx(
    markdown_text: str,
    output_path: str | pathlib.Path,
    *,
    asset_dir: str | pathlib.Path | None = None,
) -> pathlib.Path:
    """Render `markdown_text` to a .docx at `output_path`.

    `asset_dir` is the directory holding staged image files (figure_N.png);
    body image embeds resolve against it. Returns the written path.
    """
    out = pathlib.Path(output_path).expanduser()
    out.parent.mkdir(parents=True, exist_ok=True)
    assets = pathlib.Path(asset_dir).expanduser() if asset_dir else None

    tokens = _md().parse(markdown_text)
    root = SyntaxTreeNode(tokens)

    doc = Document()
    _apply_report_template(doc)
    for node in root.children:
        _render_block(doc, node, assets)
    _fix_drawing_ids(doc)  # non-zero visual ids so Hancom renders images (EXP-2)
    doc.save(str(out))
    return out
